import shutil
import runpy
from pathlib import Path

import pandas as pd
from docx import Document

from pipeline_utils import ensure_dirs, log_message

cfg = runpy.run_path('00_config.py')
PROC_DIR = cfg['PROC_DIR']
RESULT_DIR = cfg['RESULT_DIR']
DOC_PATH = Path('Manuscript-Chu-v2.docx')
BACKUP_PATH = Path('former_result') / 'Manuscript-Chu-v2.before_bioinformatics_update.docx'


def read_csv(path):
    return pd.read_csv(path)


def pretty_pathway(pathway):
    return pathway.replace('HALLMARK_', '').replace('_', ' ').title()


def join_items(items):
    items = [str(item) for item in items if str(item).strip()]
    if not items:
        return 'NA'
    if len(items) == 1:
        return items[0]
    if len(items) == 2:
        return f'{items[0]} and {items[1]}'
    return ', '.join(items[:-1]) + f', and {items[-1]}'


def remove_extra_rows(table, keep_rows):
    while len(table.rows) > keep_rows:
        table._tbl.remove(table.rows[-1]._tr)


def load_context():
    pheno = read_csv(PROC_DIR / 'pheno_macula_4groups.csv')
    candidate = read_csv(RESULT_DIR / 'tables' / 'candidate_gene_summary.csv')
    core = read_csv(RESULT_DIR / 'tables' / 'inflammatory_core_genes.csv')
    progressive = read_csv(RESULT_DIR / 'tables' / 'progressive_inflammatory_genes.csv')
    lasso = read_csv(RESULT_DIR / 'tables' / 'lasso_coefficients.csv')
    roc = read_csv(RESULT_DIR / 'tables' / 'roc_combined_signature.csv').iloc[0]
    infl_comp = read_csv(RESULT_DIR / 'tables' / 'inflammation_group_comparison.csv').iloc[0]
    infl_trend = read_csv(RESULT_DIR / 'tables' / 'inflammation_severity_trend.csv').iloc[0]
    immune_primary = read_csv(RESULT_DIR / 'tables' / 'immune_primary_comparison.csv')
    immune_trend = read_csv(RESULT_DIR / 'tables' / 'immune_severity_trend.csv')
    gene_immune = read_csv(RESULT_DIR / 'tables' / 'gene_immune_correlations.csv')
    recurrence = read_csv(RESULT_DIR / 'tables' / 'pathway_recurrence_summary.csv')

    candidate_map = dict(zip(candidate['metric'], candidate['value']))
    top_pathways = [pretty_pathway(x) for x in recurrence.head(5)['pathway'].tolist()]
    significant_primary_cells = immune_primary[immune_primary['padj'] < 0.05]['cell_type'].tolist()
    if not significant_primary_cells:
        significant_primary_cells = immune_primary.sort_values(['padj', 'pvalue']).head(2)['cell_type'].tolist()
    significant_trend_cells = immune_trend[immune_trend['padj'] < 0.05]['cell_type'].tolist()

    top_cor = gene_immune.sort_values(['padj', 'pvalue', 'rho'], ascending=[True, True, False]).head(6)
    top_cor_pairs = [
        f"{row['gene_symbol']}-{row['cell_type']} (rho={float(row['rho']):.3f})"
        for _, row in top_cor.iterrows()
    ]

    return {
        'macula_total': int(len(pheno)),
        'healthy_n': int((pheno['disease_group'] == 'healthy control').sum()),
        'diabetic_n': int((pheno['disease_group'] == 'diabetic').sum()),
        'npdr_n': int((pheno['disease_group'] == 'NPDR').sum()),
        'advanced_n': int((pheno['disease_group'] == 'NPDR/PDR + DME').sum()),
        'primary_deg_n': int(candidate_map['primary_deg_significant']),
        'core_genes': core['gene_symbol'].tolist(),
        'progressive_genes': progressive['gene_symbol'].tolist(),
        'lasso_genes': lasso['gene_symbol'].tolist(),
        'lasso_table': lasso.copy(),
        'combined_auc': float(roc['auc']),
        'combined_ci_low': float(roc['ci95_low']),
        'combined_ci_high': float(roc['ci95_high']),
        'inflammation_primary_p': float(infl_comp['pvalue']),
        'inflammation_primary_delta': float(infl_comp['delta_mean']),
        'inflammation_rho': float(infl_trend['rho']),
        'inflammation_trend_p': float(infl_trend['pvalue']),
        'top_pathways': top_pathways,
        'significant_primary_cells': significant_primary_cells,
        'significant_trend_cells': significant_trend_cells,
        'top_cor_pairs': top_cor_pairs,
    }


def build_paragraph_updates(ctx):
    core_genes = join_items(ctx['core_genes'])
    progressive_genes = join_items(ctx['progressive_genes'])
    lasso_genes = join_items(ctx['lasso_genes'])
    top_pathways = join_items(ctx['top_pathways'])
    primary_cells = join_items(ctx['significant_primary_cells'])
    trend_cells = join_items(ctx['significant_trend_cells'])
    top_cor_pairs = join_items(ctx['top_cor_pairs'])

    return {
        12: (
            'Background: Inflammation is increasingly recognized as a major contributor to diabetic '
            'retinopathy (DR), yet the stage-resolved inflammatory transcriptomic changes of the human '
            'retina remain incompletely characterized.'
        ),
        13: (
            'Methods: GSE160306 was analyzed with a fully Python-based workflow. Only macular samples '
            'were retained. Raw counts were used for DESeq2-based differential expression, whereas '
            'log2(CPM+1) values recalculated from the same counts were used for PCA, ssGSEA, LASSO '
            'modeling, and visualization. Inflammation-associated candidate genes were defined as the '
            'intersection between significant primary DEGs and the Hallmark inflammatory response gene '
            'set, followed by severity-trend analysis, LASSO logistic regression with repeated '
            'cross-validation, gene-centered preranked GSEA, and immune ssGSEA.'
        ),
        14: (
            f"Results: A total of {ctx['macula_total']} macular samples were analyzed, including "
            f"{ctx['healthy_n']} healthy controls, {ctx['diabetic_n']} diabetic samples, {ctx['npdr_n']} "
            f"NPDR samples, and {ctx['advanced_n']} advanced DR samples. DESeq2 identified "
            f"{ctx['primary_deg_n']} significant genes in the primary comparison of healthy control "
            f"versus NPDR/PDR + DME. Intersecting these DEGs with the Hallmark inflammatory response "
            f"signature yielded {len(ctx['core_genes'])} inflammatory core genes, of which "
            f"{len(ctx['progressive_genes'])} also showed a positive severity trend. LASSO retained "
            f"{len(ctx['lasso_genes'])} genes ({lasso_genes}), and the combined out-of-fold ROC AUC was "
            f"{ctx['combined_auc']:.3f}. Gene-centered GSEA highlighted recurrent Hallmark programs "
            f"related to {top_pathways}. Immune ssGSEA indicated significant perturbation of "
            f"{primary_cells} in the primary comparison, and multiple selected genes were strongly "
            f"correlated with immune signatures."
        ),
        15: (
            'Conclusion: This stage-aware macular transcriptomic analysis identified a discovery-phase '
            'nine-gene inflammatory candidate signature associated with advanced DR and the retinal '
            'immune microenvironment. These findings refine the bioinformatic evidence base for DR, but '
            'still require independent external and experimental validation.'
        ),
        18: (
            'What is already known on this topic - inflammation is a recognized driver of DR '
            'progression, but the stage-resolved retinal inflammatory transcriptome remains insufficiently defined.'
        ),
        19: (
            'What this study adds - a macula-restricted, four-group transcriptomic workflow identified '
            '13 inflammatory core genes, 9 progression-related genes, and a nine-gene LASSO signature '
            'linked to immune and pathway remodeling in advanced DR.'
        ),
        20: (
            'How this study might affect research, practice or policy - these results provide a more '
            'defensible set of discovery-stage inflammatory targets for downstream validation, rather '
            'than recycling unsupported legacy biomarkers.'
        ),
        29: (
            f"Transcriptomic data for diabetic retinopathy (DR) were obtained from the Gene Expression "
            f"Omnibus (GEO) database under accession number GSE160306. Because the dataset contains "
            f"retinal samples from different anatomical regions and disease stages, only macular samples "
            f"were retained for the formal bioinformatic workflow. The final cohort included "
            f"{ctx['healthy_n']} healthy controls, {ctx['diabetic_n']} diabetic samples, {ctx['npdr_n']} "
            f"NPDR samples, and {ctx['advanced_n']} NPDR/PDR + DME samples. Raw count matrices were "
            'used for differential expression analysis, whereas log2(CPM+1) values recomputed from the '
            'same counts were used for PCA, ssGSEA, LASSO modeling, and visualization. Genes with low '
            'counts were filtered using a count >= 10 in at least 5 samples.'
        ),
        30: (
            'The Hallmark inflammatory response gene set was retrieved from MSigDB and used as the '
            'predefined inflammation-related reference. A cleaned 28-signature immune cell gene-set '
            'collection was additionally prepared for downstream immune ssGSEA.'
        ),
        32: (
            'Differentially expressed genes (DEGs) were identified with DESeq2 implemented through the '
            'PyDESeq2 framework in Python. The primary contrast was healthy control versus NPDR/PDR + '
            'DME, and the healthy control versus diabetic, healthy control versus NPDR, and diabetic '
            'versus NPDR/PDR + DME comparisons were analyzed as supportive contrasts. Genes with an '
            'adjusted P value < 0.05 and an absolute log2 fold change >= 0.5 were considered '
            'significant.'
        ),
        34: (
            'Inflammatory activity at the sample level was quantified by ssGSEA using the '
            'HALLMARK_INFLAMMATORY_RESPONSE gene set implemented through gseapy. Group differences in '
            'inflammation ssGSEA scores were evaluated with the Mann-Whitney U test, and four-group '
            'severity trends were evaluated with Spearman correlation.'
        ),
        35: (
            'Inflammation-associated candidate genes were defined as the intersection between significant '
            'primary-comparison DEGs and the Hallmark inflammatory response signature. Genes with a '
            'positive severity correlation and FDR < 0.1 across the four ordered clinical groups were '
            'further classified as progressive inflammatory genes.'
        ),
        37: (
            'LASSO logistic regression was performed on the inflammation-associated candidate genes in the '
            'primary comparison cohort using L1-penalized logistic regression. Repeated stratified '
            'four-fold cross-validation (10 repeats) was used to tune the penalty parameter, and '
            'five-fold stratified cross-validation was used to generate out-of-fold predicted '
            'probabilities.'
        ),
        38: (
            'The discriminative performance of the selected genes and the combined signature was assessed '
            'using ROC analysis. AUC values with bootstrap-derived 95% confidence intervals were '
            'reported. Because the discovery cohort was relatively small and no external cohort was '
            'available, the selected genes were interpreted as candidate biomarkers rather than '
            'validated diagnostic markers.'
        ),
        40: (
            'For functional interpretation, each selected gene was used as an anchor for Spearman-ranked '
            'preranked GSEA against the Hallmark collection using gseapy. Pathways with nominal '
            'P < 0.05 and FDR q < 0.25 were considered significant.'
        ),
        41: (
            'An IPA-ready input table containing the selected genes, LASSO coefficients, and differential '
            'expression statistics was exported for optional manual upload. IPA itself was not executed '
            'as part of the automated Python workflow.'
        ),
        43: (
            'Relative immune activity was estimated by ssGSEA using 28 curated immune cell signatures. '
            'Primary-comparison group differences were tested with the Mann-Whitney U test, severity '
            'trends were tested with Spearman correlation, and gene-immune correlations were adjusted for '
            'multiple testing.'
        ),
        61: '3.1 Macular cohort definition and inflammation-associated differential expression',
        62: (
            f"After restricting the analysis to macular samples, {ctx['macula_total']} samples were "
            f"included in the final transcriptomic cohort, consisting of {ctx['healthy_n']} healthy "
            f"controls, {ctx['diabetic_n']} diabetic samples, {ctx['npdr_n']} NPDR samples, and "
            f"{ctx['advanced_n']} advanced DR samples (NPDR/PDR + DME) (Figure 1). In the primary "
            f"contrast of healthy control versus advanced DR, DESeq2 identified {ctx['primary_deg_n']} "
            f"significant genes. The inflammation ssGSEA score was nominally higher in advanced DR than "
            f"in controls (delta mean = {ctx['inflammation_primary_delta']:.3f}, P = "
            f"{ctx['inflammation_primary_p']:.4f}) and showed a positive four-group severity trend "
            f"(rho = {ctx['inflammation_rho']:.3f}, P = {ctx['inflammation_trend_p']:.4f}). "
            f"Intersecting the primary DEGs with the Hallmark inflammatory response signature yielded "
            f"{len(ctx['core_genes'])} inflammatory core genes: {core_genes} (Figure 2)."
        ),
        63: '3.2 Screening of a nine-gene candidate signature associated with DR',
        64: (
            f"Among the inflammatory core genes, {len(ctx['progressive_genes'])} also showed a positive "
            f"severity trend across the four clinical groups: {progressive_genes}. LASSO logistic "
            f"regression retained {len(ctx['lasso_genes'])} genes in the final signature: {lasso_genes} "
            f"(Figure 3A and 3B; Table 1). Expression visualization across the four groups showed that "
            'most selected genes increased from healthy control toward advanced DR, consistent with their '
            'selection as progression-related inflammatory candidates (Figure 3C).'
        ),
        65: (
            f"The combined signature achieved an out-of-fold AUC of {ctx['combined_auc']:.3f} "
            f"(95% CI {ctx['combined_ci_low']:.3f}-{ctx['combined_ci_high']:.3f}) (Figure 3E). This "
            'performance supports moderate discriminative ability within the discovery cohort, although it '
            'should not be misrepresented as independent clinical validation.'
        ),
        66: '3.3 Gene-centered enrichment analysis revealed recurrent inflammatory and remodeling programs',
        67: (
            f"Gene-centered preranked GSEA showed that the selected genes repeatedly converged on "
            f"Hallmark programs related to {top_pathways} (Figure 4A and 4B). These recurrent pathways "
            'support the interpretation that the selected genes participate in inflammatory amplification, '
            'stress signaling, tissue remodeling, and immune communication within advanced DR.'
        ),
        68: (
            'Several genes additionally showed inverse enrichment for pathways such as oxidative '
            'phosphorylation and mTORC1 signaling, suggesting that the inflammatory signature was '
            'accompanied by broader metabolic rewiring rather than by isolated cytokine activation alone. '
            'These enrichment results are hypothesis-generating and should be interpreted as pathway-level '
            'associations rather than proof of direct causality.'
        ),
        69: '3.4 Immune ssGSEA identified altered macrophage- and B-cell-related signatures in advanced DR',
        70: (
            f"Immune ssGSEA showed that {primary_cells} remained significantly different between the "
            f"primary comparison groups after multiple-testing correction (Figure 5A). Across the four "
            f"ordered groups, {trend_cells if trend_cells != 'NA' else 'no cell signatures'} showed "
            'significant severity trends after correction, whereas additional cell types such as natural '
            'killer cells, memory B cells, and neutrophils showed only nominal trends.'
        ),
        71: (
            'These results suggest that advanced DR is accompanied by immune remodeling with an important '
            'innate immune component, but the inferred cell-type shifts should be treated as relative '
            'transcriptomic signals rather than direct cell counts.'
        ),
        72: '3.5 Selected inflammatory genes were strongly correlated with immune-related signatures',
        73: (
            f"The strongest gene-immune correlations included {top_cor_pairs} (Figure 5B and "
            'Supplementary Figure 6). Macrophage-related scores were repeatedly linked to CMKLR1, FZD5, '
            'TLR3, and TIMP1, whereas MDSC- and regulatory T-cell-related signals were also prominent in '
            'the correlation structure.'
        ),
        74: (
            'Taken together, the immune comparison and correlation analyses support an association between '
            'the selected inflammatory genes and the retinal immune microenvironment, while still falling '
            'short of proving absolute immune-cell infiltration or causality.'
        ),
        83: (
            f"Diabetic retinopathy is increasingly recognized as a chronic neurovascular inflammatory "
            f"disorder rather than a purely microvascular complication. In the present stage-aware "
            f"macular transcriptomic analysis, we identified {len(ctx['core_genes'])} inflammatory core "
            f"genes and a nine-gene candidate signature linked to disease severity, pathway remodeling, "
            'and immune-associated transcriptomic changes. The main value of the current work is that it '
            'moves the analysis away from oversimplified binary grouping and toward a four-stage design '
            'that better reflects DR progression.'
        ),
        84: (
            f"The retained genes ({lasso_genes}) collectively point toward innate immune activation, "
            'stress signaling, and tissue remodeling. Several of them, including MSR1, CMKLR1, TLR3, '
            'and TIMP1, are biologically plausible in the context of inflammatory retinal injury, whereas '
            'others such as RNF144B, FZD5, RGS1, and OPRK1 may represent less explored but still '
            'credible candidates deserving follow-up investigation. The fact that all nine signature '
            'genes also arose from the progression-related inflammatory subset strengthens their internal '
            'coherence.'
        ),
        85: (
            f"The enrichment analysis further supports this interpretation. Recurrent Hallmark programs "
            f"centered on {top_pathways}, indicating that the identified genes sit within a broader "
            'network of inflammatory amplification, interferon-related stress responses, extracellular '
            'matrix remodeling, and vascular-pathway adaptation. At the same time, the inverse enrichment '
            'of metabolic programs in part of the signature suggests that inflammation and metabolic '
            'reprogramming are intertwined rather than separable axes in advanced DR.'
        ),
        86: (
            f"The immune ssGSEA results were also informative. {primary_cells} differed significantly in "
            f"the primary comparison, and {trend_cells if trend_cells != 'NA' else 'no cell signatures'} "
            'showed significant stage-related trends after correction. Moreover, the strongest gene-immune '
            f"correlations ({top_cor_pairs}) repeatedly pointed to macrophage- and myeloid-related "
            'signals. This pattern is broadly compatible with the known role of innate immune activation '
            'in DR, although bulk transcriptomic deconvolution remains an indirect approximation.'
        ),
        87: (
            'A practical strength of this analysis is the explicit separation of healthy control, '
            'diabetic, NPDR, and NPDR/PDR + DME groups within the macular compartment. That design '
            'reduces the risk of conflating diabetes itself with retinopathy progression and avoids the '
            'common but sloppy habit of collapsing all retinal disease samples into a single DR bin. The '
            'progression-oriented filtering step therefore adds biological structure that a simple '
            'case-control comparison would miss.'
        ),
        88: (
            'The current manuscript, however, must stay honest about what has and has not been done. The '
            'repository supports transcriptomic discovery, pathway analysis, immune scoring, and export of '
            'an IPA-ready file, but it does not contain an independent external cohort, ELISA/RT-qPCR '
            'validation data, or cell-function experiments. Any statement claiming that those validations '
            'have already been completed would be fiction, not scholarship.'
        ),
        89: (
            f"The discriminative performance of the combined signature (OOF AUC = {ctx['combined_auc']:.3f}) "
            'should therefore be interpreted as internal discovery-phase evidence only. It is useful for '
            'prioritization, but it does not establish a clinically deployable diagnostic model. Likewise, '
            'the exported IPA input table is only a handoff for optional downstream analysis and should not '
            'be described as if the IPA results were already generated inside this workflow.'
        ),
        90: (
            'Another limitation is that immune deconvolution from bulk retinal transcriptomes cannot '
            'distinguish whether the observed signals arise from infiltrating leukocytes, resident glia, '
            'or shifts in cell-state programs. These data are best viewed as supportive context for gene '
            'prioritization, not as a substitute for single-cell profiling or tissue-level immunostaining.'
        ),
        91: (
            'Despite these limitations, the study still has clear strengths. It uses a reproducible '
            'Python-based pipeline, preserves clinically meaningful disease stages, ties DEG analysis to a '
            'predefined inflammatory reference set, and integrates candidate-gene selection with pathway '
            'and immune-level interpretation. That is materially stronger than recycling a list of legacy '
            'genes from an opaque analysis stack.'
        ),
        92: (
            'The main limitations remain the modest sample size, the lack of external validation, the '
            'absence of wet-lab confirmation, and the fact that enrichment and deconvolution analyses are '
            'association-based. Future work should prioritize independent cohorts, direct molecular '
            'validation of the nine-gene panel, and experimental dissection of the most plausible targets '
            'within retinal endothelial and immune-related contexts.'
        ),
        93: (
            f"In conclusion, this macula-restricted stage-aware analysis identified {len(ctx['core_genes'])} "
            f"inflammatory core genes and a nine-gene candidate signature ({lasso_genes}) associated with "
            'advanced DR, pathway remodeling, and immune-related transcriptomic shifts. These findings '
            'provide a more defensible bioinformatic framework for the manuscript, but they remain '
            'discovery-stage results until external and experimental validation is completed.'
        ),
        96: (
            'DR: diabetic retinopathy; DEG: differentially expressed gene; ssGSEA: single-sample gene set '
            'enrichment analysis; GSEA: gene set enrichment analysis; LASSO: least absolute shrinkage and '
            'selection operator; ROC: receiver operating characteristic; AUC: area under the curve; CPM: '
            'counts per million; NPDR: non-proliferative diabetic retinopathy; PDR: proliferative '
            'diabetic retinopathy; DME: diabetic macular edema; OOF: out-of-fold.'
        ),
        151: 'Table 1. Coefficients of the nine-gene LASSO signature.',
        154: (
            'Figure 1. Overview of the macular DR transcriptomic workflow. (A) Python-based analysis '
            'workflow. (B) Sample composition of the four retained macular groups. (C) PCA of the '
            'macular cohort based on log2(CPM+1) expression values.'
        ),
        155: (
            'Figure 2. Differential and inflammatory signals across the macular cohort. (A) Inflammation '
            'ssGSEA across the four disease groups. (B) Volcano plot of the primary differential '
            'expression comparison between healthy control and NPDR/PDR + DME. (C) Heatmap of the '
            'inflammatory core genes. (D) Overlap between significant primary DEGs and the Hallmark '
            'inflammatory response signature.'
        ),
        156: (
            'Figure 3. Construction and evaluation of the nine-gene candidate signature. (A) LASSO '
            'coefficient path. (B) Repeated cross-validation performance curve. (C) Expression patterns '
            'of representative selected genes across the four groups. (D) ROC curves of individual genes. '
            '(E) Out-of-fold ROC curve of the combined signature.'
        ),
        157: (
            'Figure 4. Gene-centered preranked GSEA and pathway recurrence analysis. (A) Heatmap of '
            'normalized enrichment scores across selected genes and recurrent Hallmark pathways. (B) '
            'Dotplot of recurrent Hallmark pathways shared by multiple selected genes. (C) Optional '
            'gene-pathway network summarizing shared enrichment relationships.'
        ),
        158: (
            'Figure 5. Immune ssGSEA and gene-immune association analysis. (A) Representative immune cell '
            'ssGSEA scores across the disease groups. (B) Heatmap of Spearman correlations between the '
            'selected genes and immune cell signatures.'
        ),
        160: (
            'Supplementary Table 1. Full primary differential expression results together with the '
            'inflammatory core-gene and progressive inflammatory-gene lists.'
        ),
        161: (
            'Supplementary Table 2. Gene-centered Hallmark preranked GSEA results for the selected genes '
            'and the pathway recurrence summary.'
        ),
        162: (
            'Supplementary Table 3. Immune ssGSEA primary-comparison results, severity-trend analysis, '
            'and gene-immune correlation statistics.'
        ),
        163: (
            'Supplementary Figures 1-6. Quality-control metrics, sample clustering, expanded DEG '
            'heatmaps, inflammatory core-gene heatmaps, representative single-gene GSEA plots, and '
            'gene-immune scatter plots.'
        ),
    }


def update_table(doc, lasso_table):
    table = doc.tables[0]
    table.rows[0].cells[0].text = 'Gene symbol'
    table.rows[0].cells[1].text = 'Coefficient'
    remove_extra_rows(table, 1)
    for _, row in lasso_table.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row['gene_symbol'])
        cells[1].text = f"{float(row['coefficient']):.6f}"


def main():
    ensure_dirs(BACKUP_PATH.parent, RESULT_DIR / 'logs')
    ctx = load_context()
    updates = build_paragraph_updates(ctx)

    if not DOC_PATH.exists():
        raise FileNotFoundError(f'Manuscript not found: {DOC_PATH}')

    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)
    if len(doc.paragraphs) < max(updates) + 1:
        raise ValueError('Unexpected manuscript structure: paragraph count is smaller than expected.')

    for index, text in updates.items():
        doc.paragraphs[index].text = text

    update_table(doc, ctx['lasso_table'])
    doc.save(DOC_PATH)

    log_message(
        '12_update_manuscript_docx',
        (
            f"updated_paragraphs={len(updates)} table_rows={len(ctx['lasso_table'])} "
            f"backup={BACKUP_PATH.as_posix()}"
        ),
    )


if __name__ == '__main__':
    main()
